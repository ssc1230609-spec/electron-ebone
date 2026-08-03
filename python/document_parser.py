# -*- coding: utf-8 -*-
"""
文献解析脚本 — 将 literature/ 目录下的 PDF/Word/Markdown 文件
统一转为纯文本，按 512 token 滑动窗口分块，输出到 literature/chunks.jsonl

用法: python document_parser.py

依赖: PyMuPDF(fitz), python-docx
输出: literature/chunks.jsonl（每行一个 JSON 对象，含 text + 来源元数据）
"""
import sys
import os
import json
import re
import glob

# 强制 UTF-8
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')
if sys.stderr.encoding != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8')

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LIT_DIR = os.path.join(SCRIPT_DIR, 'literature')
CHUNKS_FILE = os.path.join(LIT_DIR, 'chunks.jsonl')

# 分块参数
TARGET_TOKENS = 512    # 每块目标 token 数
STRIDE_TOKENS = 256    # 滑动步长（50% 重叠）
CHARS_PER_TOKEN = 1.5  # 中文约 1-2 字符/token，取平均值


def emit(progress, message):
    """输出进度 JSON（复用项目统一协议）"""
    print(json.dumps({"progress": progress, "message": message}, ensure_ascii=False), flush=True)


def estimate_tokens(text):
    """粗略估算 token 数：中文字符按 1 token/字，英文按 1 token/4字符"""
    cn_chars = len(re.findall(r'[一-鿿]', text))
    other_chars = len(text) - cn_chars
    return cn_chars + other_chars // 4


def clean_text(text):
    """清洗提取文本：去控制字符、私用区(PUA)占位符、规整空白。
    fitz 对无 Unicode 映射的 CID 字体会返回 U+E000–U+F8FF 私用区占位符，属于噪声，需剔除。"""
    # C0 控制字符（保留 \t \n）
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
    # Unicode 私用区字符（无法映射的字体占位符）
    text = re.sub(r'[-]', '', text)
    # 折叠连续空格/制表符
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # 3+ 换行折叠为 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ============================================================
# 文本提取
# ============================================================
def extract_pdf(filepath):
    """从 PDF 提取文本（用 PyMuPDF/fitz，正确处理中文 CID 字体 PDF）"""
    try:
        import fitz
        doc = fitz.open(filepath)
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text().strip()
            if text:
                pages.append({"page": i + 1, "text": text})
        doc.close()
        return pages
    except ImportError:
        emit(0, "错误: 缺少 PyMuPDF，请运行 pip install PyMuPDF")
        return []
    except Exception as e:
        emit(0, f"PDF解析失败 {filepath}: {e}")
        return []


def extract_docx(filepath):
    """从 Word 文档提取文本"""
    try:
        from docx import Document
        doc = Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)
        return [{"page": 1, "text": '\n'.join(full_text)}] if full_text else []
    except ImportError:
        emit(0, "错误: 缺少 python-docx，请运行 pip install python-docx")
        return []
    except Exception as e:
        emit(0, f"Word解析失败 {filepath}: {e}")
        return []


def extract_markdown(filepath):
    """从 Markdown 提取纯文本（去除格式标记）"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()
        # 去除 Markdown 格式
        text = re.sub(r'```[\s\S]*?```', '', text)          # 代码块
        text = re.sub(r'`[^`]+`', '', text)                  # 行内代码
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)          # 图片
        text = re.sub(r'\[([^\]]+)\]\(.*?\)', r'\1', text)   # 链接保留文字
        text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)  # 标题标记
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)      # 加粗
        text = re.sub(r'\*([^*]+)\*', r'\1', text)           # 斜体
        text = re.sub(r'^[-*+]\s', '', text, flags=re.MULTILINE)  # 列表标记
        text = re.sub(r'^\d+\.\s', '', text, flags=re.MULTILINE)  # 有序列表
        text = re.sub(r'^\s*[-=]+\s*$', '', text, flags=re.MULTILINE)  # 分隔线
        text = re.sub(r'\n{3,}', '\n\n', text)               # 多余空行
        text = text.strip()
        return [{"page": 1, "text": text}] if text else []
    except Exception as e:
        emit(0, f"Markdown解析失败 {filepath}: {e}")
        return []


# ============================================================
# 文本分块
# ============================================================
def chunk_text(text, target_tokens=TARGET_TOKENS, stride_tokens=STRIDE_TOKENS):
    """
    滑动窗口分块：按段落边界切分，保留上下文重叠
    返回文本块列表
    """
    # 按段落分割
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    if not paragraphs:
        return []

    chunks = []
    current_chunk = []
    current_tokens = 0

    for para in paragraphs:
        para_tokens = estimate_tokens(para)

        # 如果单个段落就超过目标，强制切分
        if para_tokens > target_tokens * 1.5:
            # 先把当前累积的块保存
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
                # 保留最后一部分作为重叠
                overlap = []
                overlap_tokens = 0
                for p in reversed(current_chunk):
                    pt = estimate_tokens(p)
                    if overlap_tokens + pt > stride_tokens:
                        break
                    overlap.insert(0, p)
                    overlap_tokens += pt
                current_chunk = overlap
                current_tokens = overlap_tokens

            # 按字符切分大段落
            chars = list(para)
            est_chars = int(target_tokens * CHARS_PER_TOKEN)
            stride_chars = int(stride_tokens * CHARS_PER_TOKEN)
            i = 0
            while i < len(chars):
                segment = ''.join(chars[i:i + est_chars])
                if segment.strip():
                    chunks.append(segment.strip())
                i += stride_chars
            continue

        # 正常累积
        if current_tokens + para_tokens > target_tokens and current_chunk:
            chunks.append('\n'.join(current_chunk))
            # 保留 stride 部分作为重叠
            overlap = []
            overlap_tokens = 0
            for p in reversed(current_chunk):
                pt = estimate_tokens(p)
                if overlap_tokens + pt > stride_tokens:
                    break
                overlap.insert(0, p)
                overlap_tokens += pt
            current_chunk = overlap
            current_tokens = overlap_tokens

        current_chunk.append(para)
        current_tokens += para_tokens

    # 最后一块
    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks


# ============================================================
# 主流程
# ============================================================
def parse_all():
    """扫描 literature/ 目录，解析所有文献并分块输出"""
    if not os.path.exists(LIT_DIR):
        emit(100, json.dumps({"success": False, "error": "literature/ 目录不存在"}))
        return

    # 扫描文件
    files = []
    for ext in ['*.pdf', '*.docx', '*.md', '*.markdown']:
        files.extend(glob.glob(os.path.join(LIT_DIR, '**', ext), recursive=True))

    if not files:
        emit(100, json.dumps({"success": False, "error": "未找到 PDF/Word/Markdown 文件"}))
        return

    all_chunks = []
    total_files = len(files)

    for fi, filepath in enumerate(files):
        filename = os.path.basename(filepath)
        rel_path = os.path.relpath(filepath, LIT_DIR)
        progress = int((fi / total_files) * 90)
        emit(progress, f"解析 [{fi+1}/{total_files}] {filename}")

        # 根据扩展名选择解析器
        ext = os.path.splitext(filename)[1].lower()
        if ext == '.pdf':
            pages = extract_pdf(filepath)
        elif ext == '.docx':
            pages = extract_docx(filepath)
        elif ext in ('.md', '.markdown'):
            pages = extract_markdown(filepath)
        else:
            continue

        # 分块
        for page_info in pages:
            page_num = page_info["page"]
            text = clean_text(page_info["text"])
            if not text:
                continue

            chunks = chunk_text(text)
            for ci, chunk_text_str in enumerate(chunks):
                all_chunks.append({
                    "source_file": rel_path,
                    "page": page_num,
                    "chunk_id": ci,
                    "text": chunk_text_str,
                    "tokens_est": estimate_tokens(chunk_text_str)
                })

    # 写入 chunks.jsonl
    with open(CHUNKS_FILE, 'w', encoding='utf-8') as f:
        for chunk in all_chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + '\n')

    result = {
        "success": True,
        "files_processed": total_files,
        "total_chunks": len(all_chunks),
        "output": CHUNKS_FILE
    }
    emit(100, json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    try:
        parse_all()
    except Exception as e:
        emit(100, json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)
